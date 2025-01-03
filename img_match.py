import os
import time
import fitz
import openai
import re
import docx
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_folder(folder_path):
    try:
        # 使用 makedirs 创建文件夹，如果不存在则创建
        os.makedirs(folder_path)
        print(f"文件夹 '{folder_path}' 创建成功")
    except OSError as e:
        print(f"创建文件夹 '{folder_path}' 失败: {e}")
        return 1


def list_files_only(directory_path):  # 读取文件列表
    if os.path.isdir(directory_path):
        file_list = [
            f
            for f in os.listdir(directory_path)
            if os.path.isfile(os.path.join(directory_path, f))
        ]
        file_list = [file for file in file_list if ".pdf" in file]
        return file_list
    else:
        return f"{directory_path} is not a directory"


def img_match(context, s_title, doc0, flag_t):
    """
    :param context: 该部分内容
    :param s_title: 主题文件夹
    :param doc0: 文件属性
    :param flag_t: 图片序号
    :return: 返回文件变量
    """

    def list_files_only_png(directory_path):  # 读取图片列表
        if os.path.isdir(directory_path):
            file_list = [
                f
                for f in os.listdir(directory_path)
                if os.path.isfile(os.path.join(directory_path, f))
            ]
            file_list = [file for file in file_list if ".png" in file]
            return file_list
        else:
            return f"{directory_path} is not a directory"

    def read_pdf(pdf_path, output_folder):

        dimlimit = 120  # 100  # 每个图像边缘的最小像素数限制
        relsize = 0  # 0.05  # 图像：图像尺寸比必须大于此值（5%）
        abssize = 0  # 2048  # 图像绝对大小限制 2 KB：如果小于此值，则忽略
        global Theme
        global open_ai_figure
        global time_sleep
        openai.api_key = open_ai_figure

        def recoverpix(doc, item):
            '''
            恢复像素
            :param doc:文件路径
            :param item:存储位置
            :return:
            '''
            xref = item[0]  # PDF 图像的 xref
            smask = item[1]  # 其 /SMask 的 xref

            # 特殊情况：存在 /SMask 或 /Mask
            if smask > 0:
                pix0 = fitz.Pixmap(doc.extract_image(xref)["image"])
                if pix0.alpha:  # 捕获异常情况
                    pix0 = fitz.Pixmap(pix0, 0)  # 删除 alpha 通道
                mask = fitz.Pixmap(doc.extract_image(smask)["image"])

                try:
                    pix = fitz.Pixmap(pix0, mask)
                except:  # 如果有问题，回退到原始基本图像
                    pix = fitz.Pixmap(doc.extract_image(xref)["image"])

                if pix0.n > 3:
                    ext = "pam"
                else:
                    ext = "png"

                return {  # 创建预期的字典
                    "ext": ext,
                    "colorspace": pix.colorspace.n,
                    "image": pix.tobytes(ext),
                }

            # 特殊情况：存在 /ColorSpace 定义
            # 为确保安全，我们将这些情况转换为 RGB PNG 图像
            if "/ColorSpace" in doc.xref_object(xref, compressed=True):
                pix = fitz.Pixmap(doc, xref)
                pix = fitz.Pixmap(fitz.csRGB, pix)
                return {  # 创建预期的字典
                    "ext": "png",
                    "colorspace": 3,
                    "image": pix.tobytes("png"),
                }
            return doc.extract_image(xref)

        def get_descriptive(con, flag_d):
            def replace_invalid_filename_chars(filename):
                # 定义在Windows中无效的字符
                invalid_chars = '<>:"/\\|?*'

                # 创建一个转换表，将无效字符替换为空格
                trans_table = str.maketrans(invalid_chars, ' ' * len(invalid_chars))

                # 使用转换表替换字符串中的无效字符
                return filename.translate(trans_table)

            messages = [
                {'role': 'system',
                 'content': f"You are an academician with the highest authority in the field of {Theme}.Now I'm going to "
                            f"give"
                            f"you a description of some pictures and you need to summarize one of them in 20 words based on "
                            f"my prompts. Here is the description of the pictures:{con}"
                            f"Note: Just describe one of the images I mentioned in 15 words，"
                            f"Don't mention the name of the picture"},
                {'role': 'user',
                 'content': f"Based on the text, summarize in 15 words the picture in the text that refers to the {flag_d} "
                            f"small of the serial number."
                 }]
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.1,
                max_tokens=30, )
            content = response.choices[0].message.content
            time.sleep(time_sleep)
            return replace_invalid_filename_chars(content)

        time.sleep(time_sleep)  # 保证程序正确运行
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        xreflist = []

        for pno in range(page_count):
            imglist = []
            figlist = []

            il = doc.get_page_images(pno)
            imglist.extend([x[0] for x in il])
            if len(il) == 0:
                continue
            else:
                page = doc[pno]
                text = page.get_text("blocks")
                pattern = re.compile(r'^fig.*', re.IGNORECASE)
                for t0 in text:
                    if pattern.search(t0[4]):
                        figlist.append(re.sub(r'\n', '', t0[4]))
                print(figlist)
            if len(figlist) == 0 or len(figlist) < len(il):
                continue
            for img in il:
                flag = 1
                try:
                    des = get_descriptive(figlist, flag)
                    flag = flag + 1
                    print(des)
                    xref = img[0]
                    if xref in xreflist:
                        continue
                    width = img[2]
                    height = img[3]
                    if min(width, height) <= dimlimit:
                        continue
                    image = recoverpix(doc, img)
                    n = image["colorspace"]
                    imgdata = image["image"]
                    if len(imgdata) <= abssize:
                        continue
                    if len(imgdata) / (width * height * n) <= relsize:
                        continue
                    imgfile = os.path.join(output_folder, des + '.' + "png")
                    fout = open(imgfile, "wb")
                    fout.write(imgdata)
                    fout.close()
                    xreflist.append(xref)

                except Exception as e:  # 未捕获到异常，程序直接报错
                    print(e)
                    flag = flag + 1
                    continue

    def text_matching(con, s_title, list0):
        """

        :param con: 单句文本
        :param s_title:小标题
        :return:匹配到的图片
        """
        global Theme
        global open_ai_figure
        openai.api_key = open_ai_figure
        list1 = []
        for i in list0:
            list1.append(re.sub(r'\.png$', '', i))
        messages = [
            {'role': 'system',
             'content': f"You are an academician w-ith the highest authority in the field of {Theme}.I'm matching images to a "
                        f"paragraph in a review paper of {s_title}, please select the one with the highest match from the "
                        f"list of images and just return the corresponding numerical number, or just let me know '-1' if "
                        f"you think none of them are relevant. "
                        f"Note: All you have to do is tell me the number, no need for any extra words."
                        f"The content of the paper is as follows:{con}"},
            {'role': 'user',
             'content': f"The list of images is as follows：{list0}"
             }]
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.1,
            max_tokens=2, )
        content = response.choices[0].message.content
        time.sleep(time_sleep)
        print(content)
        print(list0)
        print(con)
        try:
            result = int(content)
        except:
            result = -1

        if result == -1:
            return -1
        else:
            return list0[result]

    def add_center_picture(self, image_path_or_stream, save_path=None, fig_flag=None):
        """
        在指定位置插入图片
        :param self: 文件变量
        :param image_path_or_stream: 图片名
        :param save_path: 图片文件夹有必要时加
        :param fig_flag: 图片序号
        :return:self
        """
        if save_path is None:
            image_path_or_stream0 = image_path_or_stream
        else:
            image_path_or_stream0 = save_path + "/" + image_path_or_stream
        if fig_flag is None:
            image_path_or_stream1 = image_path_or_stream
        else:
            image_path_or_stream1 = "Fig" + str(fig_flag) + "." + image_path_or_stream
        img = Image.open(image_path_or_stream0)
        width, height = img.size
        tab = self.add_table(rows=2, cols=1)  # 添加一个2行1列的空表
        cell = tab.cell(0, 0)  # 获取某单元格对象（从0开始索引）
        ph = cell.paragraphs[0]
        ph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = ph.add_run()
        ff = max(width / 300, height / 300)
        run.add_picture(image_path_or_stream0, width=docx.shared.Pt(width / ff), height=docx.shared.Pt(height / ff))
        # 设置第二行单元格的文本
        tab.cell(1, 0).text = re.sub(r'\.png$', '', image_path_or_stream1)
        # 设置第二行单元格的段落对齐方式为居中
        ph_second_row = tab.cell(1, 0).paragraphs[0]
        ph_second_row.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return self

    global Theme
    sentences = re.split(r'[.!?]', context)
    sentences.pop()
    for i in range(len(sentences)):
        sentences[i] = sentences[i] + "."
    print(sentences)
    title0 = re.sub(r'^\d+(\.\d+)*\s*', '', s_title)
    paper_list = list_files_only(fr"data\{Theme}\{title0}")
    sentences = re.split(r'[.!?]', context)
    sentences.pop()
    for i in range(len(sentences)):
        sentences[i] = sentences[i] + "."
    dic_con = {}
    for i in sentences:
        dic_con[i] = {
            "title": "",
            "title_png": ""
        }

    flag_list = []
    for i in range(len(paper_list)):
        if title0 + "_" + str(i + 1) in context:
            flag_list.append(i + 1)
            i0 = title0 + '_' + str(i + 1)
            file_is = create_folder(fr"data\{Theme}\{title0}\{i0}")
            if file_is:
                print(55101)
                break
            read_pdf(fr"data\{Theme}\{title0}\{paper_list[i]}",
                     fr"data\{Theme}\{title0}\{i0}")  # 全部图片已被提取
    for i in flag_list:
        i0 = f"[{title0}_{str(i)}]"
        i1 = f"{title0}_{str(i)}"
        print(i0)
        sentences_with_tag = [s for s in sentences if i0 in s and s.strip()]
        print(sentences_with_tag)
        list0 = list_files_only_png(fr"data\{Theme}\{title0}\{i1}")
        for sentence in sentences_with_tag:
            dic_con[sentence]["title_png"] = text_matching(sentence, title0, list0)
            dic_con[sentence]["title"] = i1
    # doc0
    pa = doc0.add_paragraph()
    for sentence in sentences:
        pa.text = pa.text + sentence
        print(sentence)
        print(dic_con[sentence])
        print(pa)
        if dic_con[sentence]["title_png"] == -1 or dic_con[sentence]["title_png"] == '':
            continue
        else:
            doc0 = add_center_picture(doc0, dic_con[sentence]["title_png"],
                                      save_path=fr'data\{Theme}\{title0}\{dic_con[sentence]["title"]}',
                                      fig_flag=flag_t)
            pa = doc0.add_paragraph()
            flag_t = flag_t + 1
    return doc0, flag_t
#
#
# Theme = 'financial engineering'
# open_ai_figure = "your api key"
# time_sleep = 30
#
# doc = docx.Document()
# con = "Financial engineering, a multidisciplinary field amalgamating mathematical finance, computer science, and economic theory, is pivotal in designing financial products and managing risks. By utilizing mathematical models and computational tools, financial engineering optimizes investment strategies and fosters innovative financial solutions, thereby enhancing decision-making processes in the financial industry.\n\nLooking ahead at future trends in financial engineering in China, the integration of technology, particularly through financial technology (FinTech), is poised to revolutionize the industry. The utilization of new technologies in finance is reshaping the development, accessibility, and delivery of financial products and services [Future Trends_7]. The exponential growth of Information and Communication Technologies (ICT) has underscored the significance of technology integration in the financial sector, transforming how individuals engage with technology and manage their daily affairs [Future Trends_9].\n\nThe phenomena of digitization, convergence, and transcoding are driving theoretical shifts in financial engineering, rendering traditional management approaches obsolete. These changes are eroding conventional boundaries and reshaping the accessibility and delivery of financial products and services, necessitating adaptation to technological and social shifts to achieve objectives effectively [Future Trends_7]. Moreover, the impact of globalization is profound, fostering interconnectedness among economies, cultures, and societies, thereby influencing the emergence of new technologies, market expansions, and international collaborations [Future Trends_9].\n\nIn essence, the amalgamation of technology and globalization is propelling transformative changes in China's financial industry, prompting the adoption of new theoretical frameworks and innovative strategies to navigate the evolving financial services landscape. Financial engineering professionals in China must remain attuned to technological advancements and global trends to harness technology integration for sustainable growth and competitiveness in the financial sector."
# doc, flag = img_match(context=con, s_title="5 Future Trends", doc0=doc, flag_t=1)
# doc.save(r'a_fe.docx')
