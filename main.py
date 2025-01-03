import json
from langchain_openai import ChatOpenAI
import docx
from langchain_openai import OpenAIEmbeddings

from img_match import img_match
from spider_final_doi import crawler
import PyPDF2
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain_community.vectorstores import Chroma
import panel as pn  # GUI
from langchain.chains import RetrievalQA
# 构建 prompt
from class_dic import dic_stru
import re
import openai
import os
import traceback
import requests
from fake_useragent import UserAgent
from urllib.parse import quote
from bs4 import BeautifulSoup
from selenium import webdriver
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.firefox.options import Options
import time
from selenium.webdriver.common.keys import Keys
import pyautogui
from selenium.webdriver.common.by import By
import os
import time
import fitz
import openai
import re
import docx
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def keep_letters(input_string):  # 去数字
    result = re.sub(r'^\d+(\.\d+)*\s*', '', input_string)
    return str(result)


def get_context(qa_interface, theme, title1):  # 给出具体内容
    title1 = keep_letters(title1)
    time.sleep(time_sleep)
    question = f'''I am writing a review on  "{theme}" and now I need to write about the 
     "{title1}". "Summarize what is said about  "{title1}".
     Note: 1. You just have to give information related to the topic and don't give me any other irrelevant information;
          2.If there is no specific mention ,you just tell me 'no.';
          3. Do not mention the words 'context' and 'text'.
          4. Paraphrased in the tone of an academic paper.
          5. Do not start with "The given context".

       '''
    # result = qa({"question": question, "chat_history": []})
    result = qa_interface(question)
    print(result['result'])
    return result['result']


def get_context0(qa_interface, theme, title1, title2):  # 给出具体内容
    title1 = keep_letters(title1)
    title20 = ""
    for key0 in title2:
        title20 = title20 + "," + keep_letters(key0)
    print(title2)
    print(title20)
    time.sleep(time_sleep)
    #    question = f'''I am writing a review on  "{theme}" and now I need to write about the
    # "{title1}". "Give me all of the info on "{str(title20)}".
    # Note: 1. You just have to give information related to the topic and don't give me any other irrelevant information;
    #      2.If there is no specific mention ,you just tell me 'no.';
    #      3. Do not mention the words 'context' and 'text'.
    #      4. Do not start with "The given context".
    #      5. Please refer to the template I provided, which returns the absolute standard json format:{str(title20)}
    #   ''' # 用第二种
    flag_g = 1
    while flag_g == 1:
        try:
            question = f"""Give me information about {title20} of {title1}?"""
            # Please provide information on healthcare based on the context provided.
            result = qa_interface.invoke(question)
            print(result['result'])
            flag_g = 0
            return result
            return result['result']
        except:
            time.sleep(time_sleep)

    # return str(result['result'])
    # return str(result['source_documents']) + str(result['result'])


def get_retriever(detected_text):
    # 0
    chunk_size = 2000  # 设置块大小2000
    chunk_overlap = 500  # 设置块重叠大小
    global time_sleep
    global llm_name
    r_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    r_page = r_splitter.create_documents([detected_text])
    directory = 'index_store'
    vector_index = FAISS.from_documents(r_page, OpenAIEmbeddings())
    time.sleep(time_sleep)
    vector_index.save_local(directory)
    retriever = vector_index.as_retriever(search_type="similarity",
                                          search_kwargs={"k": 4})
    qa_interface = RetrievalQA.from_chain_type(
        llm=ChatOpenAI(model_name=llm_name, temperature=0.5),
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=True)
    return qa_interface


def create_folder(folder_path):
    try:
        # 使用 makedirs 创建文件夹，如果不存在则创建
        os.makedirs(folder_path)
        print(f"文件夹 '{folder_path}' 创建成功")
    except OSError as e:
        print(f"创建文件夹 '{folder_path}' 失败: {e}")
        return 1


def list_files_only(directory_path):  # 读取文件列表
    if os.path.isdir(directory_path):
        file_list = [
            f
            for f in os.listdir(directory_path)
            if os.path.isfile(os.path.join(directory_path, f))
        ]
        file_list = [file for file in file_list if ".pdf" in file]
        return file_list
    else:
        return f"{directory_path} is not a directory"


def extract_numbers(input_string):
    numbers = re.findall(r'\d+', input_string)
    return [int(number) for number in numbers]  # 返回标题级数


def get_valuation0(part_s, context):  # 评估
    global Theme
    global open_ai_stru
    openai.api_key = open_ai_stru

    messages = [
        {'role': 'system',
         'content': f"You're a senior scientist working on {Theme}."},
        {'role': 'user',
         'content': f'''The following is the specific content of the subheading on {part_s}. 
         Please assess the degree of relevance of this content to the title from an academic perspective. 
         If the degree of relevance is high, please return '2’. 
         If the degree of relevance is moderate, please return '1’. 
         If the degree of relevance is low, please return '0’. 
         If the degree of relevance is very low, please return '-1’. 
         Please note that you only need to provide the corresponding number 
         without any additional description. The details are as follows: {context}'''
         }]
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages,
        temperature=0.1,
        max_tokens=3000, )
    content = response.choices[0].message.content
    # print(response.text)
    try:
        # print(content)
        print(f"grade:{content}")
        # if int(result0) > 0:
        #     print(part_s, context)
        return int(content)  # 返回评估分数
    except:
        return 0


def get_structure0(theme):  # 获取框架
    global open_ai_stru
    global Title
    openai.api_key = open_ai_stru

    json_ex = {
        '1. Introduction': "",
        '2. Historical of Mathematics': {
            '2.1 Ancient Mathematics': {'2.1.1 Mathematics in the Middle Ages': "",
                                        '2.1.2 Renaissance ': ""},
            '2.2 Mathematics in the 20th Century': ""},
        '3. Branches of Mathematics': {'3.1 Algebra': "", '3.2 Geometry': "", '3.3 Calculus': "", '3.4 Statistics': "",
                                       '3.5 Number Theory': "", '3.6 Mathematical Logic': ""},
        '4. Conclusion': ''
    }
    messages = [
        {'role': 'system',
         'content': f"You're a senior scientist working on {theme}."},
        {'role': 'user',
         'content': f'''"I am currently working on a review paper on the topic of '{Title}’.
          To structure your essay effectively, please refer to the general essay format. and assign appropriate titles for each level of your essay. 
          Please adhere to the following guidelines:
              1.Each heading at different levels should be numbered.
              2.Only provide titles for the levels, without any additional explanatory text.
              3.The number of chapters should not exceed 6.
              4.No need for a separate title.
              5.You can use up to four levels of headings.
              6.Do not include any text other than the title and serial number.
              7.No extra line breaks
              8.Please output the result as json format, e.g:{json_ex}
              9.No more than 4 words per title
          Thank you for your cooperation!"
          '''
         }]
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages,
        temperature=0.5,
        max_tokens=3000, )
    content = response.choices[0].message.content
    # print(content)
    # structure_list0 = get_structure_list0(content)
    # dic = get_dic0(content)
    print(content)
    return content
    # return [structure_list0, dic]  # 返回输出框架和字典内容


def get_dic(qa_interface, theme, title0, dic0, Dic, flag):  # 获取内容  flag是参考文献编码
    """
    :param qa_interface: 检索器
    :param theme: 主题
    :param title0: 二级标题
    :param dic0: 三级标题
    :param Dic: 结果字典
    :param flag: 参考文献序号
    :return: 结果字典
    """
    result_dic = get_context0(qa_interface, theme, title0, dic0)
    try:
        if get_valuation0(title0, result_dic['result']) <= 1:
            print('error')
            # Dic[flag] = result_dic['source_documents']
            return Dic
        else:
            Dic[flag] = result_dic['result']
        return Dic
    except:
        print('error')
        return Dic


def get_str(qa_interface, theme, title0, Dic, flag):  # 获取内容
    result_str = get_context(qa_interface, theme, title0)
    if get_valuation0(title0, result_str) <= 1:
        print('error')
    else:
        # print(t5)
        Dic[title0] += result_str + "\n"
    return Dic


def auto_get_title(specialty, grade):
    global open_ai_stru
    global time_sleep
    global Theme
    global llm_name
    openai.api_key = open_ai_stru
    flag_auto = ''

    def ask_question(messages, question_number):
        print(f"请回答第{question_number}个问题")
        flag_auto = input("Please choose A or B:  ")
        print()
        time.sleep(time_sleep)  # 必要的延时
        messages.append({'role': 'user', 'content': flag_auto})
        return get_response(messages)

    def get_response(messages):
        response = openai.chat.completions.create(
            model=llm_name,
            messages=messages,
            max_tokens=100,
            temperature=0.3,
        )
        all_text = ''
        # for event in response:
        #     event_text = event['choices'][0]['delta']
        #     answer = event_text.get('content', '')
        #     all_text += answer
        all_text = response.choices[0].message.content
        print('Expert:', all_text)
        messages.append({'role': 'assistant', 'content': all_text})
        return all_text

    if grade == "":
        grade = "Senior undergraduate majoring"
    print("You're " + grade + " of " + specialty)
    prompt = f"""If I am a {grade} in {specialty}, and now I want to write a paper in the field of {specialty},
    please ask me 10 logical questions with two options. in the field to induce me to give an "A" or "B" answer.   
    When I finish answering the 10 questions, you will give me 3 definite choices based on my interests and directions.   
    Do you understand what I mean? Now please ask your first question and wait for me to answer."""

    messages = [{"role": "system", "content": f"You're an expert in the field of {specialty}."},
                {'role': 'user', 'content': prompt}]

    initial_response = get_response(messages)

    for i in range(1, 11):
        if i == 10:
            print(f"请回答第{i}个问题")
            flag_auto = input("Please choose A or B:  ") + " Now you can give me 3 definite choices."
        else:
            initial_response = ask_question(messages, i)
    time.sleep(time_sleep)  # 必要的延时
    messages.append({'role': 'user', 'content': flag_auto})
    final_response = get_response(messages)

    flag_auto = input("Please give me a number 1~3:  ")
    time.sleep(time_sleep)  # 必要的延时
    print()
    flag_auto = (
        f'I choose NO.{flag_auto}, Now you can give me the json format like: {{"title": "", "describe for it": ""}}. '
        f'Just give me the json statement, no other prompts are needed.')
    messages.append({'role': 'user', 'content': flag_auto})
    json_response = get_response(messages)
    return json_response  # 输出JSON格式的选题建议


def get_meta_student(part_s, context, json0):  # 评估
    global Theme
    global Title
    global open_ai_master
    openai.api_key = open_ai_master
    context = str(context)
    for key in json0.keys():
        print(key)
        key0 = keep_letters(key)
        messages = [
            {'role': 'system',
             'content': f'You are an student of {Theme} at Harvard University, your grades are the highest in the program, '
                        f'and you are writing an overview paper on the topic of "{Title}". '
                        f'I will tell you a few information about financial engineering, '
                        f'you can only write about "{key0} " in {part_s} based on the given references. '
                        f'You will need to give the appropriate reference number in the text by using [] where you cite the literature.'},
            {'role': 'user',
             'content': context
             }]

        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.2,
            max_tokens=4000, )
        content = response.choices[0].message.content
        print(1111111111111111)
        print(content)
        print(1111111111111111)
        time.sleep(time_sleep)
        # print(response.text)
        json0[key] = content
        # print(json0)
    try:
        return json0
    except:
        return json0


def get_meta_expert(part_s, context):  # 评估
    global Theme
    global Title
    global open_ai_master
    openai.api_key = open_ai_master
    for key in context.keys():
        print(context[key])
        messages = [
            {'role': 'system',
             'content': f'You are an academician with the highest authority in the field of {Theme}.As your student, I am '
                        f'writing a review paper on "{Title}". I am handing over the written part to you '
                        f'and please refer to the template I provided.'
                        f"Note: Don't change the arguments mentioned in the paper, just optimize the language structure and wording."},
            {'role': 'user',
             'content': context[key]
             }]
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.1,
            max_tokens=4000, )
        content = response.choices[0].message.content
        context[key] = content
        time.sleep(time_sleep)
    # print(response.text)
    try:
        # print(context)

        return context
    except:
        return 0


def img_match(context, s_title, doc0, flag_t):
    """
    :param context: 该部分内容
    :param s_title: 主题文件夹
    :param doc0: 文件属性
    :param flag_t: 图片序号
    :return: 返回文件变量
    """

    def list_files_only_png(directory_path):  # 读取图片列表
        if os.path.isdir(directory_path):
            file_list = [
                f
                for f in os.listdir(directory_path)
                if os.path.isfile(os.path.join(directory_path, f))
            ]
            file_list = [file for file in file_list if ".png" in file]
            return file_list
        else:
            return f"{directory_path} is not a directory"

    def read_pdf(pdf_path, output_folder):

        dimlimit = 120  # 100  # 每个图像边缘的最小像素数限制
        relsize = 0  # 0.05  # 图像：图像尺寸比必须大于此值（5%）
        abssize = 0  # 2048  # 图像绝对大小限制 2 KB：如果小于此值，则忽略
        global Theme
        global open_ai_figure
        global time_sleep
        openai.api_key = open_ai_figure

        def recoverpix(doc, item):
            '''
            恢复像素
            :param doc:文件路径
            :param item:存储位置
            :return:
            '''
            xref = item[0]  # PDF 图像的 xref
            smask = item[1]  # 其 /SMask 的 xref

            # 特殊情况：存在 /SMask 或 /Mask
            if smask > 0:
                pix0 = fitz.Pixmap(doc.extract_image(xref)["image"])
                if pix0.alpha:  # 捕获异常情况
                    pix0 = fitz.Pixmap(pix0, 0)  # 删除 alpha 通道
                mask = fitz.Pixmap(doc.extract_image(smask)["image"])

                try:
                    pix = fitz.Pixmap(pix0, mask)
                except:  # 如果有问题，回退到原始基本图像
                    pix = fitz.Pixmap(doc.extract_image(xref)["image"])

                if pix0.n > 3:
                    ext = "pam"
                else:
                    ext = "png"

                return {  # 创建预期的字典
                    "ext": ext,
                    "colorspace": pix.colorspace.n,
                    "image": pix.tobytes(ext),
                }

            # 特殊情况：存在 /ColorSpace 定义
            # 为确保安全，我们将这些情况转换为 RGB PNG 图像
            if "/ColorSpace" in doc.xref_object(xref, compressed=True):
                pix = fitz.Pixmap(doc, xref)
                pix = fitz.Pixmap(fitz.csRGB, pix)
                return {  # 创建预期的字典
                    "ext": "png",
                    "colorspace": 3,
                    "image": pix.tobytes("png"),
                }
            return doc.extract_image(xref)

        def get_descriptive(con, flag_d):
            global open_ai_figure
            openai.api_key = open_ai_figure

            def replace_invalid_filename_chars(filename):
                # 定义在Windows中无效的字符
                invalid_chars = '<>:"/\\|?*'

                # 创建一个转换表，将无效字符替换为空格
                trans_table = str.maketrans(invalid_chars, ' ' * len(invalid_chars))

                # 使用转换表替换字符串中的无效字符
                return filename.translate(trans_table)

            messages = [
                {'role': 'system',
                 'content': f"You are an academician with the highest authority in the field of {Theme}.Now I'm going to "
                            f"give"
                            f"you a description of some pictures and you need to summarize one of them in 20 words based on "
                            f"my prompts. Here is the description of the pictures:{con}"
                            f"Note: Just describe one of the images I mentioned in 15 words，"
                            f"Don't mention the name of the picture"},
                {'role': 'user',
                 'content': f"Based on the text, summarize in 15 words the picture in the text that refers to the {flag_d} "
                            f"small of the serial number."
                 }]
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.1,
                max_tokens=30, )
            content = response.choices[0].message.content
            time.sleep(time_sleep)
            return replace_invalid_filename_chars(content)

        time.sleep(time_sleep)  # 保证程序正确运行
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        xreflist = []

        for pno in range(page_count):
            imglist = []
            figlist = []

            il = doc.get_page_images(pno)
            imglist.extend([x[0] for x in il])
            if len(il) == 0:
                continue
            else:
                page = doc[pno]
                text = page.get_text("blocks")
                pattern = re.compile(r'^fig.*', re.IGNORECASE)
                for t0 in text:
                    if pattern.search(t0[4]):
                        figlist.append(re.sub(r'\n', '', t0[4]))
                print(figlist)
            if len(figlist) == 0 or len(figlist) < len(il):
                continue
            for img in il:
                flag = 1
                try:
                    des = get_descriptive(figlist, flag)
                    flag = flag + 1
                    print(des)
                    xref = img[0]
                    if xref in xreflist:
                        continue
                    width = img[2]
                    height = img[3]
                    if min(width, height) <= dimlimit:
                        continue
                    image = recoverpix(doc, img)
                    n = image["colorspace"]
                    imgdata = image["image"]
                    if len(imgdata) <= abssize:
                        continue
                    if len(imgdata) / (width * height * n) <= relsize:
                        continue
                    imgfile = os.path.join(output_folder, des + '.' + "png")
                    fout = open(imgfile, "wb")
                    fout.write(imgdata)
                    fout.close()
                    xreflist.append(xref)

                except Exception as e:  # 未捕获到异常，程序直接报错
                    print(e)
                    flag = flag + 1
                    continue

    def text_matching(con, s_title, list0):
        """

        :param con: 单句文本
        :param s_title:小标题
        :return:匹配到的图片
        """
        global Theme
        global open_ai_figure
        openai.api_key = open_ai_figure
        list1 = []
        for i in list0:
            list1.append(re.sub(r'\.png$', '', i))
        messages = [
            {'role': 'system',
             'content': f"You are an academician w-ith the highest authority in the field of {Theme}.I'm matching images to a "
                        f"paragraph in a review paper of {s_title}, please select the one with the highest match from the "
                        f"list of images and just return the corresponding numerical number, or just let me know '-1' if "
                        f"you think none of them are relevant. "
                        f"Note: All you have to do is tell me the number, no need for any extra words."
                        f"The content of the paper is as follows:{con}"},
            {'role': 'user',
             'content': f"The list of images is as follows：{list0}"
             }]
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.1,
            max_tokens=2, )
        content = response.choices[0].message.content
        time.sleep(time_sleep)
        print(content)
        print(list0)
        print(con)
        try:
            result = int(content)
        except:
            result = -1

        if result == -1:
            return -1
        else:
            return list0[result]

    def add_center_picture(self, image_path_or_stream, save_path=None, fig_flag=None):
        """
        在指定位置插入图片
        :param self: 文件变量
        :param image_path_or_stream: 图片名
        :param save_path: 图片文件夹有必要时加
        :param fig_flag: 图片序号
        :return:self
        """
        if save_path is None:
            image_path_or_stream0 = image_path_or_stream
        else:
            image_path_or_stream0 = save_path + "/" + image_path_or_stream
        if fig_flag is None:
            image_path_or_stream1 = image_path_or_stream
        else:
            image_path_or_stream1 = "Fig" + str(fig_flag) + "." + image_path_or_stream
        img = Image.open(image_path_or_stream0)
        width, height = img.size
        tab = self.add_table(rows=2, cols=1)  # 添加一个2行1列的空表
        cell = tab.cell(0, 0)  # 获取某单元格对象（从0开始索引）
        ph = cell.paragraphs[0]
        ph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = ph.add_run()
        ff = max(width / 300, height / 300)
        run.add_picture(image_path_or_stream0, width=docx.shared.Pt(width / ff), height=docx.shared.Pt(height / ff))
        # 设置第二行单元格的文本
        tab.cell(1, 0).text = re.sub(r'\.png$', '', image_path_or_stream1)
        # 设置第二行单元格的段落对齐方式为居中
        ph_second_row = tab.cell(1, 0).paragraphs[0]
        ph_second_row.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return self

    global Theme
    sentences = re.split(r'[.!?]', context)
    sentences.pop()
    for i in range(len(sentences)):
        sentences[i] = sentences[i] + "."
    print(sentences)
    title0 = re.sub(r'^\d+(\.\d+)*\s*', '', s_title)
    paper_list = list_files_only(fr"data\{Theme}\{title0}")
    sentences = re.split(r'[.!?]', context)
    sentences.pop()
    for i in range(len(sentences)):
        sentences[i] = sentences[i] + "."
    dic_con = {}
    for i in sentences:
        dic_con[i] = {
            "title": "",
            "title_png": ""
        }

    flag_list = []
    for i in range(len(paper_list)):
        if title0 + "_" + str(i + 1) in context:
            flag_list.append(i + 1)
            i0 = title0 + '_' + str(i + 1)
            file_is = create_folder(fr"data\{Theme}\{title0}\{i0}")
            if file_is:
                break
            read_pdf(fr"data\{Theme}\{title0}\{paper_list[i]}",
                     fr"data\{Theme}\{title0}\{i0}")  # 全部图片已被提取
    for i in flag_list:
        i0 = f"[{title0}_{str(i)}]"
        i1 = f"{title0}_{str(i)}"
        print(i0)
        sentences_with_tag = [s for s in sentences if i0 in s and s.strip()]
        print(sentences_with_tag)
        list0 = list_files_only_png(fr"data\{Theme}\{title0}\{i1}")
        for sentence in sentences_with_tag:
            dic_con[sentence]["title_png"] = text_matching(sentence, title0, list0)
            dic_con[sentence]["title"] = i1
    # doc0
    pa = doc0.add_paragraph()
    for sentence in sentences:
        pa.text = pa.text + sentence
        print(sentence)
        print(dic_con[sentence])
        print(pa)
        if dic_con[sentence]["title_png"] == -1 or dic_con[sentence]["title_png"] == '':
            continue
        else:
            doc0 = add_center_picture(doc0, dic_con[sentence]["title_png"],
                                      save_path=fr'data\{Theme}\{title0}\{dic_con[sentence]["title"]}',
                                      fig_flag=flag_t)
            pa = doc0.add_paragraph()
            flag_t = flag_t + 1
    return doc0, flag_t


def main():
    # List5: 存放各级标题
    # Dic  : 存放每节内容
    global time_sleep
    global Theme
    global Title

    flag_fig = 1  # 图片序号
    Theme = input("What's your specialty?   ")
    grade0 = input("What's your grade? If you are a senior undergraduate majoring, just input enter.  ")
    # print(time_sleep)
    choose = input("Do you have a clear research topic?\n\bYou need choose 1:YES or 2:NO.")
    if choose == "1":
        Title = input("Your topic is: ")
    else:
        print("Now I will help you choose you topic.")
        try:
            Title = auto_get_title(Theme, grade0)
        except:
            print("I'm sorry. Something went wrong.")
            Title = input("Please determine your own title: ")

    # 创建保存的文档
    doc = docx.Document()
    doc.add_heading(Title)
    doc.save(f'{Theme}_output.docx')

    Dic = json.loads(get_structure0(Theme))
    main_dic = dic_stru(Dic)
    # 检索文献并下载
    for title in main_dic.dic_keys:
        print(title)
        crawler(Theme, re.sub(r'^\d+(\.\d+)*\s*', '', title))
    for title in main_dic.str_keys:
        print(title)
        crawler(Theme, re.sub(r'^\d+(\.\d+)*\s*', '', title))
    # 获取内容
    for title, dic0 in zip(main_dic.dic_keys, main_dic.dic_list):
        print(title, dic0)
        dic1 = {}  # 用于存储
        # 创建一个 PyPDFLoader Class 实例，输入为待加载的pdf文档路径
        title0 = re.sub(r'^\d+(\.\d+)*\s*', '', title)  # 子标题
        paper_list = list_files_only(f"data\{Theme}\{title0}")
        print(paper_list)
        # 提取参考文献
        flag0 = 1
        for paper in paper_list:
            detected_text = ""
            try:  # 读取文献
                # 调用 PyPDFLoader Class 的函数 load对pdf文件进行加载
                pdf_file_obj = open(f"data\{Theme}\{title0}\\" + paper, "rb")
                pdf_reader = PyPDF2.PdfReader(pdf_file_obj)
                num_pages = len(pdf_reader.pages)
                for page_num in range(num_pages):
                    page_obj = pdf_reader.pages[page_num]
                    detected_text += page_obj.extract_text() + "\n\n"
                pdf_file_obj.close()
                print(len(detected_text))  # 找出字符数
                qa = get_retriever(detected_text)
            except:
                traceback.print_exc()
                print(f" File {paper} is empty")
                continue
            # print(qa.invoke(f"What's the main idea of {title0}?")['result'])
            dic1 = get_dic(qa, Theme, title, dic0, dic1, keep_letters(title) + "_" + str(flag0))
            flag0 = flag0 + 1
        print(dic1)
        dic_s = get_meta_student(title0, dic1, dic0)  # 学生写的
        print(dic_s)
        dic_e = get_meta_expert(title0, dic_s)  # 专家写的
        print(dic_e)
        for key in dic_e:
            main_dic.all_dic[key] = dic_e[key]

    for title in main_dic.str_keys:
        print(title)
        # 创建一个 PyPDFLoader Class 实例，输入为待加载的pdf文档路径
        title0 = re.sub(r'^\d+(\.\d+)*\s*', '', title)
        paper_list = list_files_only(f"data\{Theme}\{title0}")
        dic0 = {title: ""}
        dic1 = {}
        print(paper_list)
        # 提取参考文献
        flag0 = 1
        for paper in paper_list:

            detected_text = ""
            try:  # 读取文献
                # 调用 PyPDFLoader Class 的函数 load对pdf文件进行加载
                pdf_file_obj = open(f"data\{Theme}\{title0}\\" + paper, "rb")
                pdf_reader = PyPDF2.PdfReader(pdf_file_obj)
                num_pages = len(pdf_reader.pages)
                for page_num in range(num_pages):
                    page_obj = pdf_reader.pages[page_num]
                    detected_text += page_obj.extract_text() + "\n\n"
                pdf_file_obj.close()
                print(len(detected_text))  # 找出字符数
                qa = get_retriever(detected_text)
            except:
                traceback.print_exc()
                print(f" File {paper} is empty")
                continue

            # print(qa.invoke(f"What's the main idea of {title0}?")['result'])
            dic1 = get_dic(qa, Theme, title, dic0, dic1, keep_letters(title) + "_" + str(flag0))
            flag0 = flag0 + 1
        dic_s = get_meta_student(title0, dic1, dic0)  # 学生写的
        dic_e = get_meta_expert(title0, dic_s)  # 专家写的
        for key in dic_e:
            main_dic.all_dic[key] = dic_e[key]
    try:
        dic_text = ''
        for dic0 in main_dic.all_dic:
            dic_text = dic_text + dic0 + "\n" + main_dic.all_dic[dic0]
        qa_interface_final = get_retriever(dic_text)
        time.sleep(time_sleep)
        main_dic.all_dic["Abstract"] = qa_interface_final.invoke("What's the main idea?")['result']  # 出问题、、
        time.sleep(time_sleep)
        main_dic.all_dic["Conclusion"] = qa_interface_final.invoke("Summarize what is given.")['result']
    except:
        print("Summarize the errors that occurred")
    print(main_dic.all_dic)
    end_dic = {"Abstract": main_dic.all_dic["Abstract"]}
    for key0 in main_dic.all_dic:
        if key0 == "Abstract":
            continue
        end_dic[key0] = main_dic.all_dic[key0]

    # 遍历字典并将内容添加到Word文档
    flag_key = ""
    for key, value in end_dic.items():
        if key in main_dic.dic_keys or key in main_dic.str_keys:
            flag_key = key
        if value == "":
            doc.add_paragraph(key)
        else:
            try:
                doc.add_paragraph(key)
                doc, flag_fig = img_match(context=value, s_title=flag_key, doc0=doc, flag_t=flag_fig)
            except:
                print("No match for figure")
                doc.add_paragraph(value)

    # 保存Word文档
    doc.save(f'{Theme}_output.docx')


if __name__ == '__main__':
    # version2024.2.12
    """
    1. 修改了get_structure0的prompt
    2. 加入了json问答
    3. 实现了自动确定选题
    """
    """
    spider_final_doi 第78、118行不能用相对路径

    """
    flag = 0
    # 全局变量 后续与要加入独立的的api
    os.environ["OPENAI_API_KEY"] = \
        'your api key'


    open_ai_master = "your api key"
    open_ai_figure = "your api key"
    open_ai_stru = "your api key"
    # llm_name = "gpt-3.5-turbo"
    llm_name = "gpt-4"
    # 导入文本分割器
    chunk_size = 2000  # 设置块大小2000
    chunk_overlap = 500  # 设置块重叠大小
    # time_sleep = 20 # 用3.5时需加延时
    if llm_name == "gpt-3.5-turbo":
        time_sleep = 30  # 用3.5时需加延时
    else:
        time_sleep = 1
    Theme = ""  # 专业
    Title = ""  # 论文题目
    main()
# 先考虑二级标题即2.2
# 字典更新
# export https_proxy=http://127.0.0.1:33210 http_proxy=http://127.0.0.1:33210 all_proxy=socks5://127.0.0.1:33211
# 写框架
"""
@inproceedings{shao2024assisting,
      title={{Assisting in Writing Wikipedia-like Articles From Scratch with Large Language Models}}, 
      author={Yijia Shao and Yucheng Jiang and Theodore A. Kanell and Peter Xu and Omar Khattab and Monica S. Lam},
      year={2024},
      booktitle={Proceedings of the 2024 Conference of the North American Chapter of the Association for Computational Linguistics: Human Language Technologies, Volume 1 (Long and Short Papers)}
}
"""
